# app.py
from flask import Flask, render_template, request, send_file
from docx import Document
import pandas as pd
import os
import re
from io import BytesIO
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

if not os.path.exists("generated_files"):
    os.makedirs("generated_files")

def sanitize_filename(text):
    return re.sub(r'[^a-zA-Z0-9_-]', '_', text)

# Load Excel data
price_data = pd.read_excel("PRICELIST.xlsx", sheet_name=None)

# STRIP leading/trailing spaces from all column names in each sheet to avoid KeyErrors
for sheet in price_data:
    price_data[sheet].columns = price_data[sheet].columns.str.strip()

# Helper function to fill Word template
def fill_template(template_path, placeholders):
    left_aligned_keys = ["{{CLIENT}}"]
    
    doc = Document(template_path)
    
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                para.text = para.text.replace(key, value.upper())
                if key in left_aligned_keys:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in placeholders.items():
                        if key in para.text:
                            para.text = para.text.replace(key, value.upper())
                            if key in left_aligned_keys:
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        proposal_type = request.form.get("proposal_type")
        format_choice = request.form.get("format")

        # Gather input values
        client_data = {
            "{{CLIENT}}": request.form.get("client").upper(),
            "{{EMAIL}}": request.form.get("email"),
            "{{CONTACT}}": request.form.get("contact"),
            "{{LOCATION}}": request.form.get("location"),
            "{{ELECTRIC_BILL}}": request.form.get("electric_bill"),
            "{{ENERGY_RATE}}": request.form.get("energy_rate"),
            "{{EXPIRY_DATE}}": request.form.get("expiry")
        }

        placeholders = dict(client_data)

        if proposal_type == "hybrid":
            size = request.form.get("system_size")
            placeholders["{{SYSTEM_SIZE}}"] = size
            df = price_data["Hybrid"]

            # Debug: print column names to check for 'SYSTEM SIZE'
            print("Hybrid sheet columns:", df.columns)

            row = df[df["SYSTEM SIZE"].astype(str) == size].iloc[0]

            # Fill Part 2
            placeholders.update({
                "{{HYBRID_PV_SIZE}}": str(row["PV SYSTEM SIZE"]),
                "{{HYBRID_AREA}}": str(row["SOLAR PANELS AREA COVERED (Sq. M.)"]),
                "{{HYBRID_MONTHLY_ENERGY}}": str(row["MONTHLY GENERATED ENERGY (kWh)"]),
                "{{HYBRID_PANELS}}": str(row["NO. OF SOLAR PANELS (545W EACH)"]),
                "{{HYBRID_INVERTER}}": str(row["INVERTER SIZE"]),
                "{{HYBRID_PREMIUM}}": str(row["TOTAL PREMIUM"]),
                "{{HYBRID_SAVINGS}}": str(row["ESTIMATED MONTHLY SAVINGS"]),
                "{{HYBRID_ROI}}": str(row["RETURN OF INVESTMENT"]),

                # Part 4
                "{{H_PREMIUM}}": str(row["PREMIUM (VAT NOT INCLUDED)"]),
                "{{H_30}}": str(row["30% RESERVATION FEE / DOWNPAYMENT BEFORE MOBILIZATION"]),
                "{{H_20}}": str(row["20% UPON MATERIALS DELIVERY ON SITE"]),
                "{{H_45}}": str(row["45% UPON COMPLETION"]),
                "{{H_5}}": str(row["5% UPON TESTING AND COMMISSIONING"]),
                "{{H_TOTAL}}": str(row["TOTAL PAYMENTS"]),
                "{{H_PREMIUM5}}": str(row["PREMIUM (VAT NOT INCLUDED) +5%"])
            })

            template_path = "HYBRID-20250521-TEMPLATE.docx"

        else:
            zero = request.form.get("zero_bill_size")
            lower = request.form.get("lower_bill_size")
            placeholders["{{ZERO_BILL_SIZE}}"] = zero
            placeholders["{{LOWER_BILL_SIZE}}"] = lower

            df = price_data["OnGrid"]

            # Debug: print column names to check for 'SYSTEM SIZE'
            print("OnGrid sheet columns:", df.columns)

            zero_row = df[df["SYSTEM SIZE"].astype(str) == zero].iloc[0]
            lower_row = df[df["SYSTEM SIZE"].astype(str) == lower].iloc[0]

            # Fill Part 2
            placeholders.update({
                # WORRY FREE
                "{{ZERO_PV_SIZE}}": str(zero_row["PV SYSTEM SIZE (kWp)"]),
                "{{ZERO_AREA}}": str(zero_row["SOLAR PANELS AREA COVERED (Sq. M.)"]),
                "{{ZERO_MONTHLY_ENERGY}}": str(zero_row["MONTHLY GENERATED ENERGY (kWh)"]),
                "{{ZERO_PANELS}}": str(zero_row["NO. OF SOLAR PANELS (580W EACH)"]),
                "{{ZERO_INVERTER}}": str(zero_row["INVERTER SIZE"]),
                "{{ZERO_PREMIUM}}": str(zero_row["PREMIUM (VAT NOT INCLUDED)"]),
                "{{ZERO_SAVINGS}}": str(zero_row["ESTIMATED MONTHLY SAVINGS"]),
                "{{ZERO_ROI}}": str(zero_row["RETURN OF INVESTMENT"]),

                # WORRY LESS
                "{{LOWER_PV_SIZE}}": str(lower_row["PV SYSTEM SIZE (kWp)"]),
                "{{LOWER_AREA}}": str(lower_row["SOLAR PANELS AREA COVERED (Sq. M.)"]),
                "{{LOWER_MONTHLY_ENERGY}}": str(lower_row["MONTHLY GENERATED ENERGY (kWh)"]),
                "{{LOWER_PANELS}}": str(lower_row["NO. OF SOLAR PANELS (580W EACH)"]),
                "{{LOWER_INVERTER}}": str(lower_row["INVERTER SIZE"]),
                "{{LOWER_PREMIUM}}": str(lower_row["PREMIUM (VAT NOT INCLUDED)"]),
                "{{LOWER_SAVINGS}}": str(lower_row["ESTIMATED MONTHLY SAVINGS"]),
                "{{LOWER_ROI}}": str(lower_row["RETURN OF INVESTMENT"]),

                # Part 4 - Worry Free
                "{{OF_PREMIUM}}": str(zero_row["PREMIUM (VAT NOT INCLUDED)"]),
                "{{OF_30}}": str(zero_row["30% RESERVATION FEE / DOWNPAYMENT BEFORE MOBILIZATION"]),
                "{{OF_20}}": str(zero_row["20% UPON MATERIALS DELIVERY ON SITE"]),
                "{{OF_45}}": str(zero_row["45% UPON COMPLETION"]),
                "{{OF_5}}": str(zero_row["5% UPON TESTING AND COMMISSIONING"]),
                "{{OF_TOTAL}}": str(zero_row["TOTAL PAYMENTS"]),
                "{{OF_PREMIUM5}}": str(zero_row["PREMIUM (VAT NOT INCLUDED) +5%"]),

                # Part 4 - Worry Less
                "{{OL_PREMIUM}}": str(lower_row["PREMIUM (VAT NOT INCLUDED)"]),
                "{{OL_30}}": str(lower_row["30% RESERVATION FEE / DOWNPAYMENT BEFORE MOBILIZATION"]),
                "{{OL_20}}": str(lower_row["20% UPON MATERIALS DELIVERY ON SITE"]),
                "{{OL_45}}": str(lower_row["45% UPON COMPLETION"]),
                "{{OL_5}}": str(lower_row["5% UPON TESTING AND COMMISSIONING"]),
                "{{OL_TOTAL}}": str(lower_row["TOTAL PAYMENTS"]),
                "{{OL_PREMIUM5}}": str(lower_row["PREMIUM (VAT NOT INCLUDED) +5%"])
            })

            template_path = "ONGRID-20250521.docx"

        # Fill document
        filled_doc = fill_template(template_path, placeholders)
        output_stream = BytesIO()
        filled_doc.save(output_stream)
        output_stream.seek(0)

        safe_client = sanitize_filename(request.form.get("client"))
        safe_type = sanitize_filename(proposal_type)
        filename = f"{safe_client}-{safe_type}.docx"
        return send_file(output_stream, as_attachment=True, download_name=filename, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    return render_template("form.html")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)

